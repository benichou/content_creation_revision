from  docx import Document
from django.conf import settings
from azure.ai.formrecognizer import DocumentAnalysisClient 
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import AzureAIDocumentIntelligenceLoader
from io import BytesIO
import tempfile


def analyze_docx(file, default_credential):

    
    document_analysis_client = DocumentAnalysisClient(
        endpoint    = settings.ENDPOINT, 
        credential  = default_credential, 
        api_version = settings.FORMRECOGNIZER_API_VERSION
        )
   
    # Ensure the file pointer is at the start
    file.seek(0)

    # Read the file stream as binary
    file_stream = file.read()

    poller = document_analysis_client.begin_analyze_document(
            "prebuilt-read", file_stream)
    result = poller.result()
        
    data = []
    for paragraph in result.paragraphs:
        data.append(paragraph.content)

    return "\n".join(data)

def analyze_pptx(file, default_credential):

    document_analysis_client = DocumentAnalysisClient(
        endpoint    = settings.ENDPOINT, 
        credential  = default_credential, 
        api_version = settings.FORMRECOGNIZER_API_VERSION
        )
   
    # Ensure the file pointer is at the starts
    file.seek(0)

    # Read the file stream as binary
    file_stream = file.read()

    poller = document_analysis_client.begin_analyze_document(
            "prebuilt-read", file_stream)
    result = poller.result()
        
    data = []
    for paragraph in result.paragraphs:
        data.append(paragraph.content)

    return "\n".join(data)

def analyze_html(file, default_credential):
    
    document_analysis_client = DocumentAnalysisClient(
        endpoint    = settings.ENDPOINT, 
        credential  = default_credential, 
        api_version = settings.FORMRECOGNIZER_API_VERSION
        )
    

    file.seek(0)
    
    file_stream = file.read()
    try:
        file_data = file_stream.decode('utf-8')
    except UnicodeDecodeError:
        file_data = file_stream.decode('iso-8859-1')
    if "<html>" not in file_data and "</html>" not in file_data:
        file_data = "<html>" + file_data + "</html>"
        poller = document_analysis_client.begin_analyze_document(
            "prebuilt-read", file_data.encode('utf-8'))
    else:
        poller = document_analysis_client.begin_analyze_document(
            "prebuilt-read", file_stream)
    
    result = poller.result()
    
    print(result.content)
    
    data = []
    

    return result.content
    
    


def analyze_pdf(file, default_credential):
    # sample document
    
    document_analysis_client = DocumentAnalysisClient(
        endpoint    = settings.ENDPOINT, 
        credential  = default_credential, 
        api_version = settings.FORMRECOGNIZER_API_VERSION
        )
   
    # Ensure the file pointer is at the start
    file.seek(0)

    # Read the file stream as binary
    file_stream = file.read()

    poller = document_analysis_client.begin_analyze_document(
            "prebuilt-document", file_stream)
    result = poller.result()
        
    data = []
    
    for page in result.pages:

        for line in page.lines:
            data.append(line.content)
    return "\n".join(data)
    

def analyze_txt(file):

    file.seek(0)

    try:
        # Try reading the file with UTF-8 encoding for English text
        text = file.read().decode('utf-8')
    except UnicodeDecodeError:
        # If a UnicodeDecodeError occurs, seek back to the start and try a different encoding
        file.seek(0)
        try:
            # Attempt to read with ISO-8859-1 (Latin-1) for French and English text
            text = file.read().decode('iso-8859-1')
        except UnicodeDecodeError:
            # Handle other potential encodings or error handling here
            print("Failed to decode the file with UTF-8 and ISO-8859-1 encodings.")
            
            raise UnicodeDecodeError
            

    return text

def analyze_vtt(file):

    file.seek(0)

    try:
        # Try reading the file with UTF-8 encoding for English text
        text = file.read().decode('utf-8')
    except UnicodeDecodeError:
        # If a UnicodeDecodeError occurs, seek back to the start and try a different encoding
        file.seek(0)
        try:
            # Attempt to read with ISO-8859-1 (Latin-1) for French and English text
            text = file.read().decode('iso-8859-1')
        except UnicodeDecodeError:
            # Handle other potential encodings or error handling here
            print("Failed to decode the file with UTF-8 and ISO-8859-1 encodings.")
            
            raise UnicodeDecodeError
            

    return text    

def analyze_xls(file, default_credential):
            
    document_analysis_client = DocumentAnalysisClient(
        endpoint    = settings.ENDPOINT, 
        credential  = default_credential, 
        api_version = settings.FORMRECOGNIZER_API_VERSION
        )
   
    # Ensure the file pointer is at the starts
    file.seek(0)

    # Read the file stream as binary
    file_stream = file.read()

    poller = document_analysis_client.begin_analyze_document(
            "prebuilt-read", file_stream)
    result = poller.result()
        
    data = []
    for paragraph in result.paragraphs:
        data.append(paragraph.content)

    return "\n".join(data)