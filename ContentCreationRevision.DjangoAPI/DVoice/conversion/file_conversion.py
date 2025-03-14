import os
from pathlib import Path
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image as PILImage
import requests
from io import BytesIO
import re

from typing import Any, Union, Dict

EXTRACTED_OUTPUT_ROOT = "extracted_output" ## need to have a config file!


def add_hyperlink(paragraph: Paragraph, 
                  text: str, 
                  url: str) -> None:
    """
    Adds a hyperlink to a paragraph in a docx document.
 
    This function creates a hyperlink within the given paragraph. The hyperlink will 
    display the provided text and will link to the specified URL.
 
    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph to which the hyperlink will be added.
        text (str): The display text of the hyperlink.
        url (str): The URL the hyperlink will point to.
    Returns:
        None: This function does not return anything; it modifies the paragraph in place.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    new_run.append(r_pr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_image(doc: Document, 
              img_src: str, 
              output_dir: str, 
              max_width: float = 6.0) -> None:
    """
    Adds an image to the docx document, resizing it to fit within the specified width.
    This function allows the addition of an image to a docx document, with the image
    being resized to a maximum width (in inches) while maintaining its aspect ratio.
 
    Args:
        doc (Document): The docx Document object to which the image will be added.
        img_src (str): The source of the image, which can be a local file path or a URL.
        output_dir (str): The directory where the resized image will be temporarily saved.
        max_width (float, optional): The maximum width of the image in inches. Defaults to 6.0.
 
    Returns:
        None: This function does not return anything, it directly modifies the docx document.
    """
    try:
        # Open image from URL or local file
        if img_src.startswith("http://") or img_src.startswith("https://"):
            response = requests.get(img_src)
            image = PILImage.open(BytesIO(response.content))
        else:
            image = PILImage.open(img_src)

        # Resize image to fit within max_width
        width, height = image.size
        aspect_ratio = height / width
        new_width = min(width, int(max_width * 96))  # Convert inches to pixels
        new_height = int(new_width * aspect_ratio)

        # Save the resized image temporarily
        temp_image_path = output_dir + "\\temp_image_resized.png"
        image = image.resize(
            (new_width, new_height), resample=PILImage.Resampling.LANCZOS
        )
        image.save(temp_image_path)
        # Add the image to the document
        doc.add_picture(temp_image_path, width=Inches(new_width / 96))
    except Exception as e:
        print(f"Error adding image: {e}")


# def combine_markdown_files(input_files, output_file):
#     """
#     Combine multiple markdown files into a single markdown file.

#     Parameters:
#     - input_files: List of paths to the input markdown files.
#     - output_file: Path to the output markdown file.
#     """
#     with open(output_file, "w", encoding="utf-8") as outfile:
#         for file_index, file_path in enumerate(input_files):
#             if not os.path.isfile(file_path):
#                 print(f"Warning: {file_path} does not exist and will be skipped.")
#                 continue

#             with open(file_path, "r", encoding="utf-8") as infile:
#                 content = infile.read()
#                 # Write content into the output file
#                 outfile.write(f"<!-- Start of {os.path.basename(file_path)} -->\n")
#                 outfile.write(content)
#                 outfile.write("\n\n")  # Add space between sections
#                 outfile.write(f"<!-- End of {os.path.basename(file_path)} -->\n")

#             if file_index < len(input_files) - 1:  # Separate files visually
#                 outfile.write("\n---\n\n")  # Markdown separator for clarity

#     print(f"Combined markdown file created at: {output_file}")


math_operator_1 = r"$\times$"  # Multiplication
math_operator_2 = r"$\cdot$"  # Dot product or multiplication
math_operator_3 = r"$\\div$"  # Division
math_operator_4 = r"$\\pm$"  # Plus or minus
math_operator_5 = r"$\\mp$"  # Minus or plus
math_operator_6 = r"$\\sqrt{}$"  # Square root
math_operator_7 = r"$\\leq$"  # Less than or equal to
math_operator_8 = r"$\\geq$"  # Greater than or equal to
math_operator_9 = r"$<$"  # Less than
math_operator_10 = r"$>$"  # Greater than
math_operator_11 = r"$=$"  # Equals
math_operator_12 = r"$\\neq$"  # Not equal
math_operator_13 = r"$\\approx$"  # Approximately
math_operator_14 = r"$\\infty$"  # Infinity
math_operator_15 = r"$\\pi$"  # Pi
math_operator_16 = r"$\\alpha$"  # Alpha
math_operator_17 = r"$\\beta$"  # Beta
math_operator_18 = r"$\\gamma$"  # Gamma
math_operator_19 = r"$\\delta$"  # Delta
math_operator_20 = r"$\\int$"  # Integral
math_operator_21 = r"$\\sum$"  # Summation
math_operator_22 = r"$\\prod$"  # Product
math_operator_23 = r"$\\sin$"  # Sine function
math_operator_24 = r"$\\cos$"  # Cosine function
math_operator_25 = r"$\\tan$"  # Tangent function
math_operator_26 = r"$\\log$"  # Logarithm
math_operator_27 = r"$\\ln$"  # Natural logarithm
math_operator_28 = r"$\\theta$"  # Theta
math_operator_29 = r"$\\lambda$"  # Lambda
math_operator_30 = r"$\\Delta$"  # Delta (uppercase)
math_operator_31 = r"$\\partial$"  # Partial derivative
math_operator_32 = r"$\\nabla$"  # Nabla (gradient)
math_operator_33 = r"$\\propto$"  # Proportional to
math_operator_34 = r"$\\forall$"  # For all
math_operator_35 = r"$\\exists$"  # There exists
math_operator_36 = r"$\\subset$"  # Subset
math_operator_37 = r"$\\supset$"  # Superset
math_operator_38 = r"$\\cup$"  # Union
math_operator_39 = r"$\\cap$"  # Intersection
math_operator_40 = r"$\\Rightarrow$"  # Implies
math_operator_41 = r"$\\Leftrightarrow$"  # If and only if
math_operator_42 = r"$\\rightarrow$"  # Right arrow
math_operator_43 = r"$\\leftarrow$"  # Left arrow
math_operator_44 = r"$\\uparrow$"  # Up arrow
math_operator_45 = r"$\\downarrow$"  # Down arrow
math_operator_46 = r"\Rightarrow$"  # Implies (FOR MISTAKE FROM LLM MARKDOWN GENERATION --> IT IS LIKELY FOR IMPLIES)
math_operator_47 = r"$\Rightarrow$"  # Implies (FOR MISTAKE FROM LLM MARKDOWN GENERATION --> IT IS LIKELY FOR IMPLIES)

# Dictionary mapping Markdown math notations to their normal math symbols
math_operators_markdown = {
    math_operator_1: "×",  # Multiplication
    math_operator_2: "·",  # Dot product or multiplication
    math_operator_3: "÷",  # Division
    math_operator_4: "±",  # Plus or minus
    math_operator_5: "∓",  # Minus or plus
    math_operator_6: "√",  # Square root
    math_operator_7: "≤",  # Less than or equal to
    math_operator_8: "≥",  # Greater than or equal to
    math_operator_9: "<",  # Less than
    math_operator_10: ">",  # Greater than
    math_operator_11: "=",  # Equals
    math_operator_12: "≠",  # Not equal
    math_operator_13: "≈",  # Approximately
    math_operator_14: "∞",  # Infinity
    math_operator_15: "π",  # Pi
    math_operator_16: "α",  # Alpha
    math_operator_17: "β",  # Beta
    math_operator_18: "γ",  # Gamma
    math_operator_19: "δ",  # Delta
    math_operator_20: "∫",  # Integral
    math_operator_21: "∑",  # Summation
    math_operator_22: "∏",  # Product
    math_operator_23: "sin",  # Sine function
    math_operator_24: "cos",  # Cosine function
    math_operator_25: "tan",  # Tangent function
    math_operator_26: "log",  # Logarithm
    math_operator_27: "ln",  # Natural logarithm
    math_operator_28: "θ",  # Theta
    math_operator_29: "λ",  # Lambda
    math_operator_30: "Δ",  # Delta (uppercase)
    math_operator_31: "∂",  # Partial derivative
    math_operator_32: "∇",  # Nabla (gradient)
    math_operator_33: "∝",  # Proportional to
    math_operator_34: "∀",  # For all
    math_operator_35: "∃",  # There exists
    math_operator_36: "⊂",  # Subset
    math_operator_37: "⊃",  # Superset
    math_operator_38: "∪",  # Union
    math_operator_39: "∩",  # Intersection
    math_operator_40: "⇒",  # Implies
    math_operator_41: "⇔",  # If and only if
    math_operator_42: "→",  # Right arrow
    math_operator_43: "←",  # Left arrow
    math_operator_44: "↑",  # Up arrow
    math_operator_45: "↓",  # Down arrow
    math_operator_46: "⇒",  # Implies
}

import re

def convert_latex_to_text(markdown_text: str) -> str:
    """
    Converts LaTeX fractions in a markdown text to a human-readable format.
    This function identifies LaTeX fractions in the form of \frac{numerator}{denominator}
    and converts them into a plain text format as "numerator / denominator".
 
    Args:
        markdown_text (str): The markdown text that may contain LaTeX fractions.
 
    Returns:
        str: The markdown text with LaTeX fractions converted to a human-readable format.
    """
    # Regular expression to match LaTeX \frac{numerator}{denominator}
    pattern = r"\\frac\{([^{}]+)\}\{([^{}]+)\}"
    # Replace LaTeX fraction with "Numerator / Denominator"
    markdown_text = re.sub(pattern, r"\1 / \2", markdown_text)
    return markdown_text



def convert_markdown_math_to_normal_notation(markdown_text: str, 
                                             math_operators_markdown: dict) -> str:
    """
    Converts Markdown math notations to normal math symbols.
 
    This function replaces Markdown math notation (e.g., `\times`, `\%`) with standard math symbols
    (e.g., `*`, `%`), and processes LaTeX formatting (e.g., `\text{...}`) into plain text. It also handles
    special cases like cleaning up specific LaTeX text annotations.
 
    Args:
        markdown_text (str): The text containing Markdown math notations.
        math_operators_markdown (dict): A dictionary mapping Markdown math operators to their normal math symbols.
 
    Returns:
        str: The input text with Markdown math notations replaced by normal math symbols.
    """
    for markdown, normal in math_operators_markdown.items():
        markdown_text = markdown_text.replace(markdown, normal)

    ## cleans "$\text{Distributable Revenue}$" to Distributable Revenue ## TO DO CAN BE MADE FOR LOOP TO COVER
    ## OTHER CONDITIONS THAN JUST TEXT
    markdown_text = re.sub(r"\$\\text\{([^}]*)\}\$", r"\1", markdown_text)
    ## transforms \text{Discount Rate} to Discount Rate
    markdown_text = re.sub(r"\\text\{([^}]*)\}", r"\1", markdown_text)
    markdown_text = re.sub(r"\\times", "*", markdown_text)  # Replace \times with *
    markdown_text = re.sub(r"\\%", "%", markdown_text)      # Replace \% with %
    markdown_text = convert_latex_to_text(markdown_text)
    
    return markdown_text


def preprocess_markdown_bullet_points(md_content: str) -> str:
    """
    Processes a markdown text by escaping manually numbered bullet points.
 
    This function identifies manually numbered bullet points (e.g., `- 1.`, `- 2.`, etc.) and escapes
    the numbering by replacing it with `\.` (e.g., `- 1\.`), so it is not interpreted as a list item
    by markdown parsers.
 
    Args:
        md_content (str): The markdown content containing bullet points to be processed.
 
    Returns:
        str: The processed markdown content with escaped manually numbered bullet points.
    """
    pattern = r"^- (\d+)\."

    # Function to escape manually numbered bullet points
    def escape_manual_numbering(match):
        number = match.group(1)  # Extract the number
        return f"- {number}\\."

    # Apply regex substitution line by line
    lines = md_content.split("\n")
    processed_lines = [re.sub(pattern, escape_manual_numbering, line) for line in lines]
    return "\n".join(processed_lines)


def markdown_to_docx(md_file: str, 
                     docx_file: str) -> Document:
    """
    Converts a markdown file to a DOCX document.
 
    This function reads a markdown file, processes its content, and converts it into a DOCX document.
    It handles various markdown elements such as headings, paragraphs, hyperlinks, images, lists, and tables.
 
    Args:
        md_file (str): Path to the markdown file to be converted.
        docx_file (str): Path where the converted DOCX document will be saved.
 
    Returns:
        Document: The generated DOCX document object.
    """
    output_dir = os.getcwd()
    # Read markdown file
    try:
        with open(md_file, "r", encoding="utf-8") as file:
            md_content = file.read()
    except:
        try:
           with open(md_file, "r", encoding="latin") as file:
            md_content = file.read()
        except Exception as e:
            print(f"Error reading markdown file as {e}")
    # print(md_content)
    md_content = md_content.replace(f"/{EXTRACTED_OUTPUT_ROOT}",f"{EXTRACTED_OUTPUT_ROOT}")

    # convert markdown math notation to normal notation
    md_content = convert_markdown_math_to_normal_notation(
        md_content, math_operators_markdown
    )
    # ensure numbered bullet points are kept in word document
    md_content = preprocess_markdown_bullet_points(md_content)
    # md_content.replace(math_operator_40, "⇒")
    # Convert markdown to HTML
    html_content = markdown.markdown(md_content, extensions=["tables", "extra"])

    # Parse the HTML
    soup = BeautifulSoup(html_content, "html.parser")

    # Create a new DOCX document
    doc = Document()

    sub_bullet_point_store = []

    for element in soup.descendants:
        
        if element.name == "hr":
            doc.add_page_break()  # Insert a page break
        
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])  # Get heading level
            doc.add_heading(element.text.strip(), level=level)
            
        elif element.name == "p":
            # Check if the paragraph contains `---` for a page break
            paragraph = doc.add_paragraph()
            for child in element.children:
                if child.name == "a":  # Hyperlink
                    add_hyperlink(paragraph, child.text.strip(), child.get("href"))
                else:  # Regular text
                    # print("TEXT")
                    # print(child.text.strip())
                    if child.text.strip() not in sub_bullet_point_store:
                        paragraph.add_run(child.text.strip())
        elif element.name == "img":  # Image
            img_src = element.get("src")
            add_image(doc, img_src, output_dir)
        elif element.name == "ol":  # Handle ordered lists
            for idx, li in enumerate(element.find_all("li"), start=1):
                doc.add_paragraph(li.get_text(strip=True), style="List Bullet")
        elif element.name == "ul":  # Bullet points
            for li in element.find_all("li"):
                if bool(
                    [x.get_text() for x in li.find_all("p")]
                ):  # if sub bullet points
                    sub_bullet_points_list = [x.get_text() for x in li.find_all("p")]
                    # print("TEXT 1")
                    # print(sub_bullet_points_list)
                    for itm in sub_bullet_points_list:
                        sub_bullet_point_store.append(itm)
                        doc.add_paragraph(itm, style="List Bullet")
                else:
                    doc.add_paragraph(li.get_text(), style="List Bullet")
        elif element.name == "table":  # table data extraction
            rows = element.find_all("tr")
            if rows:
                table = doc.add_table(
                    rows=len(rows), cols=len(rows[0].find_all(["td", "th"]))
                )
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.text.strip()

    # Save the DOCX document
    print(f"saving docx file at {docx_file}")

    doc.save(docx_file)
    
    return doc

def convert_to_markdown_and_docx_document(file_path: str, 
                                          file_content: str) -> tuple:
    """
    Converts content to both a markdown (.md) file and a DOCX document.
 
    This function takes file content and converts it to a markdown file, then
    generates a DOCX document from the markdown file. The output files are saved
    in a structured directory format based on the input file name.
 
    Args:
        file_path (str): The path of the input file (used for naming output files).
        file_content (str): The content to be written to the markdown and DOCX files.
 
    Returns:
        tuple: A tuple containing:
            - final_md_file_path (str): The path to the generated markdown file.
            - final_docx_path (str): The path to the generated DOCX file.
            - final_output_file_name_wt_ext (str): The final output file name with extension.
            - doc (Document): The generated DOCX document object.
    """
    formatted_file_name = Path(file_path).stem
    final_output_file_name = f"{formatted_file_name}_DVoice_final_output"
    final_output_file_name_wt_ext = f"{final_output_file_name}.docx"
    # for page in ordered_extracted_slide_content:
    #     for page_idx, extracted_content in page.items():
            
    #         md_path = str(md_ind_pages_from_pptx) + f"\\{formatted_file_name}_page_{page_idx}.md"
            
    #         with open(md_path, "w", encoding="utf-8") as markdown_file:
    #             markdown_file.write(extracted_content["content_summary"])
    #             md_path_repo.append(md_path)
    out_path = Path(os.getcwd() + f"\\DVoice\\{EXTRACTED_OUTPUT_ROOT}\\")
    os.makedirs(out_path, exist_ok=True)
    md_out_path = Path(os.getcwd() + f"\\DVoice\\{EXTRACTED_OUTPUT_ROOT}\\markdown\\{formatted_file_name}\\")
    os.makedirs(md_out_path, exist_ok=True)
    doc_out_path = Path(os.getcwd() + f"\\DVoice\\{EXTRACTED_OUTPUT_ROOT}\\docx\\{formatted_file_name}")
    os.makedirs(doc_out_path, exist_ok=True)
    final_md_file_path = str(md_out_path) + f"\\{final_output_file_name}.md"
    with open(final_md_file_path, "w", encoding="utf-8") as markdown_file:
        markdown_file.write(file_content)
    # Create a unique markdown file
    # combine_markdown_files(md_path_repo, final_md_file_path)
    print(f"Converted to final markdown file at:{final_md_file_path}")
    # Create a unique docx file
    final_docx_path = str(doc_out_path) + f"\\{final_output_file_name_wt_ext}"
    doc = markdown_to_docx(final_md_file_path, final_docx_path)
    
    return final_md_file_path, final_docx_path, final_output_file_name_wt_ext, doc
